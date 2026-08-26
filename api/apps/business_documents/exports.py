#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from __future__ import annotations

import hashlib
import html
import io
import re
from datetime import UTC, datetime
from typing import Any
from urllib.parse import urlsplit

from docx import Document as WordDocument

from api.apps.business_documents.assets import published_template, rendering_policy
from api.apps.business_documents.errors import BusinessDocumentError, ConflictError, ValidationError
from api.db.db_models import BusinessDocument, BusinessDocumentExportArtifact, BusinessDocumentJob, BusinessDocumentRevision
from common.misc_utils import get_uuid
from common.time_utils import current_timestamp


_FORMAT_META = {
    "MARKDOWN": ("md", "text/markdown; charset=utf-8"),
    "DOCX": ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "EVA_WIKI": ("html", "text/html; charset=utf-8"),
}


def _hash_bytes(content: bytes) -> str:
    return f"sha256:{hashlib.sha256(content).hexdigest()}"


def _safe_filename(value: str) -> str:
    normalized = re.sub(r"[^\w.-]+", "_", value.strip(), flags=re.UNICODE).strip("._")
    return normalized[:160] or "business_requirements"


def _safe_external_url(value: str) -> str:
    parsed = urlsplit(value)
    if parsed.scheme.lower() not in {"http", "https"} or not parsed.netloc:
        raise ValidationError("UNSAFE_EXPORT_URL", "Only absolute HTTP(S) links are allowed in exported documents")
    return value


def _artifact_dict(row: BusinessDocumentExportArtifact) -> dict[str, Any]:
    revision = BusinessDocumentRevision.get_or_none(BusinessDocumentRevision.id == row.revision_id)
    return {
        "artifact_id": row.id,
        "document_id": row.document_id,
        "revision_id": row.revision_id,
        "revision_number": revision.revision_number if revision is not None else None,
        "format": row.export_format,
        "filename": row.filename,
        "mime_type": row.mime_type,
        "size": row.size,
        "content_hash": row.content_hash,
        "create_time": row.create_time,
    }


class BusinessDocumentExportService:
    @classmethod
    def generate(cls, job: BusinessDocumentJob, storage=None) -> dict[str, Any]:
        if job.job_type != "GENERATE_EXPORT":
            raise ValidationError("INVALID_EXPORT_JOB", "Job is not an export request")
        document = BusinessDocument.get_or_none((BusinessDocument.id == job.document_id) & (BusinessDocument.tenant_id == job.tenant_id))
        if document is None:
            raise BusinessDocumentError("DOCUMENT_NOT_FOUND", "Business document not found", 404)
        command_payload = job.payload.get("command_payload", {})
        revision_id = command_payload.get("revision_id")
        export_format = command_payload.get("format")
        if document.lifecycle_state != "AGREED" or document.current_revision_id != revision_id:
            raise ConflictError("AGREED_REVISION_REQUIRED", "Export job no longer targets the current agreed revision")
        if document.state_version != job.source_state_version:
            raise ConflictError("STALE_AI_RESULT", "Export job targets an outdated document state")
        if export_format not in _FORMAT_META:
            raise ValidationError("INVALID_EXPORT_FORMAT", "Only MARKDOWN, DOCX and EVA_WIKI are supported")
        revision = BusinessDocumentRevision.get_or_none((BusinessDocumentRevision.id == revision_id) & (BusinessDocumentRevision.document_id == document.id))
        if revision is None:
            raise BusinessDocumentError("REVISION_NOT_FOUND", "Business document revision not found", 404)

        storage_impl = storage or cls._default_storage()
        existing = BusinessDocumentExportArtifact.get_or_none(
            (BusinessDocumentExportArtifact.document_id == document.id) & (BusinessDocumentExportArtifact.revision_id == revision.id) & (BusinessDocumentExportArtifact.export_format == export_format)
        )
        if existing is not None:
            try:
                existing_content = storage_impl.get(existing.storage_bucket, existing.storage_key)
            except Exception:
                existing_content = None
            if isinstance(existing_content, bytes) and _hash_bytes(existing_content) == existing.content_hash:
                return _artifact_dict(existing)

        content = cls._render(export_format, revision)
        extension, mime_type = _FORMAT_META[export_format]
        artifact_id = get_uuid()
        filename = f"{_safe_filename(document.title)}_r{revision.revision_number}.{extension}"
        bucket = f"{document.tenant_id}-business-documents"
        storage_key = f"exports/{document.id}/{revision.id}/{artifact_id}.{extension}"
        try:
            storage_impl.put(bucket, storage_key, content)
            stored_content = storage_impl.get(bucket, storage_key)
        except Exception as exc:
            try:
                storage_impl.rm(bucket, storage_key)
            except Exception:
                pass
            raise BusinessDocumentError("EXPORT_STORAGE_WRITE_FAILED", "Export artifact could not be durably stored", 500) from exc
        if not isinstance(stored_content, bytes) or _hash_bytes(stored_content) != _hash_bytes(content):
            try:
                storage_impl.rm(bucket, storage_key)
            except Exception:
                pass
            raise BusinessDocumentError("EXPORT_STORAGE_WRITE_FAILED", "Export artifact could not be durably stored", 500)
        timestamp = current_timestamp()
        now = datetime.now()
        database = BusinessDocumentExportArtifact._meta.database
        try:
            with database.atomic():
                if existing is not None:
                    deleted = (
                        BusinessDocumentExportArtifact.delete()
                        .where(
                            (BusinessDocumentExportArtifact.id == existing.id)
                            & (BusinessDocumentExportArtifact.content_hash == existing.content_hash)
                            & (BusinessDocumentExportArtifact.storage_key == existing.storage_key)
                        )
                        .execute()
                    )
                    if deleted != 1:
                        raise ConflictError("EXPORT_ARTIFACT_CHANGED", "Export artifact changed while it was being repaired")
                BusinessDocumentExportArtifact.create(
                    id=artifact_id,
                    document_id=document.id,
                    tenant_id=document.tenant_id,
                    owner_id=document.owner_id,
                    revision_id=revision.id,
                    export_format=export_format,
                    filename=filename,
                    mime_type=mime_type,
                    size=len(content),
                    content_hash=_hash_bytes(content),
                    storage_bucket=bucket,
                    storage_key=storage_key,
                    create_time=timestamp,
                    create_date=now,
                    update_time=timestamp,
                    update_date=now,
                )
        except Exception:
            try:
                storage_impl.rm(bucket, storage_key)
            except Exception:
                pass
            raise
        if existing is not None and (existing.storage_bucket, existing.storage_key) != (bucket, storage_key):
            try:
                storage_impl.rm(existing.storage_bucket, existing.storage_key)
            except Exception:
                pass
        return _artifact_dict(BusinessDocumentExportArtifact.get_by_id(artifact_id))

    @classmethod
    def list_artifacts(cls, tenant_id: str, actor_id: str, document_id: str) -> list[dict[str, Any]]:
        if not BusinessDocument.select().where((BusinessDocument.id == document_id) & (BusinessDocument.tenant_id == tenant_id) & (BusinessDocument.owner_id == actor_id)).exists():
            raise BusinessDocumentError("DOCUMENT_NOT_FOUND", "Business document not found", 404)
        rows = (
            BusinessDocumentExportArtifact.select()
            .where((BusinessDocumentExportArtifact.document_id == document_id) & (BusinessDocumentExportArtifact.tenant_id == tenant_id) & (BusinessDocumentExportArtifact.owner_id == actor_id))
            .order_by(BusinessDocumentExportArtifact.create_time.desc())
        )
        return [_artifact_dict(row) for row in rows]

    @classmethod
    def download(cls, tenant_id: str, actor_id: str, document_id: str, artifact_id: str, storage=None):
        artifact = BusinessDocumentExportArtifact.get_or_none(
            (BusinessDocumentExportArtifact.id == artifact_id)
            & (BusinessDocumentExportArtifact.document_id == document_id)
            & (BusinessDocumentExportArtifact.tenant_id == tenant_id)
            & (BusinessDocumentExportArtifact.owner_id == actor_id)
        )
        if artifact is None:
            raise BusinessDocumentError("EXPORT_NOT_FOUND", "Export artifact not found", 404)
        storage_impl = storage or cls._default_storage()
        try:
            content = storage_impl.get(artifact.storage_bucket, artifact.storage_key)
        except Exception as exc:
            raise BusinessDocumentError("EXPORT_STORAGE_CORRUPT", "Stored export artifact is missing or corrupt", 500) from exc
        if not isinstance(content, bytes) or _hash_bytes(content) != artifact.content_hash:
            raise BusinessDocumentError("EXPORT_STORAGE_CORRUPT", "Stored export artifact is missing or corrupt", 500)
        return _artifact_dict(artifact), content

    @staticmethod
    def _default_storage():
        from common import settings

        if settings.STORAGE_IMPL is None:
            raise RuntimeError("RAGFlow storage is not initialized")
        return settings.STORAGE_IMPL

    @classmethod
    def _render(cls, export_format: str, revision: BusinessDocumentRevision) -> bytes:
        if export_format == "MARKDOWN":
            return revision.body_markdown.encode("utf-8")
        if export_format == "DOCX":
            return cls._render_docx(revision.document_ast)
        if export_format == "EVA_WIKI":
            return cls._render_eva_wiki(revision.document_ast, revision.revision_number).encode("utf-8")
        raise ValidationError("INVALID_EXPORT_FORMAT", "Unsupported export format")

    @staticmethod
    def _render_docx(document_ast: dict[str, Any]) -> bytes:
        document = WordDocument()
        for section in document_ast["sections"]:
            level = min(9, 1 + section["id"].count("."))
            document.add_heading(f"{section['id']}. {section['title']}", level=level)
            for block in section["blocks"]:
                block_type = block["type"]
                if block_type == "paragraph":
                    document.add_paragraph(block["text"])
                elif block_type == "list":
                    for item in block["items"]:
                        document.add_paragraph(str(item), style="List Bullet")
                elif block_type == "table":
                    table = document.add_table(rows=1, cols=len(block["headers"]))
                    for index, header in enumerate(block["headers"]):
                        table.rows[0].cells[index].text = str(header)
                    for row in block["rows"]:
                        cells = table.add_row().cells
                        for index, value in enumerate(row[: len(cells)]):
                            cells[index].text = "" if value is None else str(value)
                elif block_type in {"plantuml", "bpmn"}:
                    document.add_paragraph(block["source"], style="No Spacing")
                elif block_type in {"image", "reference"}:
                    document.add_paragraph(f"{block.get('alt') or block.get('label')}: {block['url']}")
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _render_eva_wiki(document_ast: dict[str, Any], revision_number: int) -> str:
        policy = rendering_policy()["eva_wiki"]
        if not policy.get("exclude_protocol"):
            raise RuntimeError("EvaWiki export policy must exclude review protocol")
        generated_at = datetime.now(UTC).date().isoformat()
        lines = [
            f'<div class="business-requirements" data-template-version="{html.escape(published_template()["template_version"])}" data-revision="{revision_number}" data-generated-at="{generated_at}">'
        ]
        for section_index, section in enumerate(document_ast["sections"]):
            heading_level = min(6, 2 + section["id"].count("."))
            section_id = html.escape(section["id"])
            section_node_id = f"br-r{revision_number}-s{section_index + 1}"
            lines.append(f'<h{heading_level} data-id="{section_node_id}">{section_id}. {html.escape(section["title"])}</h{heading_level}>')
            for block_index, block in enumerate(section["blocks"]):
                block_type = block["type"]
                block_node_id = f"{section_node_id}-b{block_index + 1}"
                if block_type == "paragraph":
                    if str(block["text"]).strip():
                        lines.append(f'<p data-id="{block_node_id}">{html.escape(block["text"])}</p>')
                elif block_type == "list":
                    items = "".join(
                        f'<li data-id="{block_node_id}-li{item_index}"><p data-id="{block_node_id}-li{item_index}-p">{html.escape(str(item))}</p></li>'
                        for item_index, item in enumerate(block["items"], start=1)
                        if str(item).strip()
                    )
                    if items:
                        lines.append(f'<ul class="{policy["root_list_class"]}" style="list-style-type: disc;" data-id="{block_node_id}">{items}</ul>')
                elif block_type == "table":
                    headers = "".join(
                        f'<th colspan="1" rowspan="1" data-x="{x}" data-y="0" data-id="{block_node_id}-h{x}"><p data-id="{block_node_id}-h{x}-p">{html.escape(str(item)) or "&#160;"}</p></th>'
                        for x, item in enumerate(block["headers"])
                    )
                    rows = "".join(
                        f'<tr data-id="{block_node_id}-r{y}">'
                        + "".join(
                            f'<td colspan="1" rowspan="1" data-x="{x}" data-y="{y}" data-id="{block_node_id}-c{x}-{y}">'
                            f'<p data-id="{block_node_id}-c{x}-{y}-p">'
                            f"{html.escape('' if value is None else str(value)) or '&#160;'}</p></td>"
                            for x, value in enumerate(row)
                        )
                        + "</tr>"
                        for y, row in enumerate(block["rows"], start=1)
                    )
                    lines.append(
                        f'<div class="{policy["table_wrapper_class"]}" data-macros="{policy["table_wrapper_macro"]}" '
                        f'data-id="{block_node_id}"><table data-id="{block_node_id}-table">'
                        f'<thead><tr data-id="{block_node_id}-r0">{headers}</tr></thead><tbody>{rows}</tbody></table></div>'
                    )
                elif block_type in {"plantuml", "bpmn"}:
                    css_class = policy["plantuml_class"] if block_type == "plantuml" else policy["xml_class"]
                    lines.append(f'<pre class="{css_class}" data-id="{block_node_id}"><code data-id="{block_node_id}-code">{html.escape(block["source"])}</code></pre>')
                elif block_type == "image":
                    image_url = _safe_external_url(block["url"])
                    lines.append(f'<img data-id="{block_node_id}" alt="{html.escape(block["alt"])}" src="{html.escape(image_url)}" />')
                elif block_type == "reference":
                    reference_url = _safe_external_url(block["url"])
                    lines.append(f'<p data-id="{block_node_id}"><a data-id="{block_node_id}-a" href="{html.escape(reference_url)}">{html.escape(block["label"])}</a></p>')
        lines.append("</div>")
        return "\n".join(lines)
